"""Generate a "signed copy" of an uploaded document with a visual signature
seal embedded at the position the signer placed it.

For PDF the seal (circle + checkmark) is drawn exactly at pos_x/pos_y on the
given page, with a small details box alongside it. For DOCX, a floating seal
icon is anchored at pos_x/pos_y of page 1 (the page signatures are placed on),
and a plain-text signature record is appended at the end of the document.
"""
from __future__ import annotations

import io
from datetime import datetime


def generate_signed_copy(
    content: bytes,
    ext: str,
    *,
    pos_x: float,
    pos_y: float,
    page_number: int,
    signer_name: str,
    timestamp: datetime,
) -> bytes:
    ext = ext.lower()
    if ext == ".pdf":
        return _stamp_pdf(content, pos_x, pos_y, page_number, signer_name, timestamp)
    if ext == ".docx":
        return _stamp_docx(content, pos_x, pos_y, page_number, signer_name, timestamp)
    raise ValueError(f"Signing is not supported for files of type '{ext}'.")


def _draw_seal(c, cx: float, cy: float, r: float = 14) -> None:
    """Draw a green circle-with-checkmark "verified" seal centered at (cx, cy)."""
    c.setFillColorRGB(0.82, 0.98, 0.90)
    c.setStrokeColorRGB(0.06, 0.48, 0.37)
    c.setLineWidth(1.6)
    c.circle(cx, cy, r, stroke=1, fill=1)

    c.setLineWidth(2)
    check = c.beginPath()
    check.moveTo(cx - r * 0.45, cy - r * 0.05)
    check.lineTo(cx - r * 0.1, cy - r * 0.45)
    check.lineTo(cx + r * 0.5, cy + r * 0.35)
    c.drawPath(check, stroke=1, fill=0)


def _stamp_pdf(content: bytes, pos_x: float, pos_y: float, page_number: int, signer_name: str, timestamp: datetime) -> bytes:
    from pypdf import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas

    reader = PdfReader(io.BytesIO(content))
    writer = PdfWriter()

    target_idx = max(0, min(page_number - 1, len(reader.pages) - 1))

    for i, page in enumerate(reader.pages):
        if i == target_idx:
            width = float(page.mediabox.width)
            height = float(page.mediabox.height)
            x = (pos_x / 100.0) * width
            y = height - (pos_y / 100.0) * height  # PDF origin is bottom-left

            buf = io.BytesIO()
            c = canvas.Canvas(buf, pagesize=(width, height))

            # Seal (circle + checkmark) at the exact position the signer clicked
            _draw_seal(c, x, y, r=14)
            c.save()
            buf.seek(0)

            overlay = PdfReader(buf)
            page.merge_page(overlay.pages[0])
        writer.add_page(page)

    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def _seal_icon_png(size: int = 64) -> io.BytesIO:
    """A small green circle-with-checkmark "verified" seal, as a transparent PNG."""
    from PIL import Image, ImageDraw

    img = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)
    pad = 3
    draw.ellipse([pad, pad, size - pad, size - pad], fill=(209, 250, 229, 255), outline=(16, 122, 95, 255), width=4)
    draw.line(
        [(size * 0.27, size * 0.52), (size * 0.45, size * 0.70), (size * 0.74, size * 0.34)],
        fill=(16, 122, 95, 255), width=5, joint="curve",
    )
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def _add_floating_seal(doc, pos_x: float, pos_y: float, size_pt: float = 26) -> None:
    """Anchor the seal icon at (pos_x%, pos_y%) of page 1, floating over the content."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    section = doc.sections[0]
    page_w = section.page_width
    page_h = section.page_height

    size_emu = int(Pt(size_pt))
    left = max(0, int(page_w * (pos_x / 100.0)) - size_emu // 2)
    top = max(0, int(page_h * (pos_y / 100.0)) - size_emu // 2)

    # New empty paragraph at the very start of the body (lands on page 1)
    new_p = OxmlElement("w:p")
    doc.element.body.insert(0, new_p)
    paragraph = Paragraph(new_p, doc)

    run = paragraph.add_run()
    run.add_picture(_seal_icon_png(), width=size_emu, height=size_emu)

    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))
    extent = inline.find(qn("wp:extent"))
    docPr = inline.find(qn("wp:docPr"))
    graphic = inline.find(qn("a:graphic"))

    anchor = OxmlElement("wp:anchor")
    for attr, val in {
        "distT": "0", "distB": "0", "distL": "0", "distR": "0",
        "simplePos": "0", "relativeHeight": "251658240",
        "behindDoc": "0", "locked": "0", "layoutInCell": "1", "allowOverlap": "1",
    }.items():
        anchor.set(attr, val)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    pos_h = OxmlElement("wp:positionH")
    pos_h.set("relativeFrom", "page")
    off_h = OxmlElement("wp:posOffset")
    off_h.text = str(left)
    pos_h.append(off_h)
    anchor.append(pos_h)

    pos_v = OxmlElement("wp:positionV")
    pos_v.set("relativeFrom", "page")
    off_v = OxmlElement("wp:posOffset")
    off_v.text = str(top)
    pos_v.append(off_v)
    anchor.append(pos_v)

    anchor.append(extent)
    anchor.append(OxmlElement("wp:wrapNone"))
    anchor.append(docPr)
    anchor.append(graphic)

    drawing.remove(inline)
    drawing.append(anchor)


def _stamp_docx(content: bytes, pos_x: float, pos_y: float, page_number: int, signer_name: str, timestamp: datetime) -> bytes:
    from docx import Document

    doc = Document(io.BytesIO(content))

    # Seal icon floating at the exact position the signer clicked (page 1)
    _add_floating_seal(doc, pos_x, pos_y)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
